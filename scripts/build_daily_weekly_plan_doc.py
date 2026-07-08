#!/usr/bin/env python3
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
CURRICULUM_PATH = ROOT / "data/thai_curriculum_master.json"
PHRASES_PATH = ROOT / "data/thai_master_phrases.json"
ASSETS_INDEX_PATH = ROOT / "data/thai_learning_assets_index.json"
OUT_PATH = ROOT / "docs/태국어_25일_일별_주차별_학습계획.docx"

PAGE_WIDTH_IN = 6.5


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_in: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_table(table, header_fill="E8EEF5") -> None:
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    run.font.name = "Malgun Gothic"
                    run._element.rPr.rFonts.set(qn("w:ascii"), "Malgun Gothic")
                    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
                    run.font.size = Pt(9)
        if row_idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 77, 120)


def add_para(doc: Document, text: str, style: str | None = None, bold=False, color=None, size=None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    return paragraph


def phrase_index(phrases: dict) -> dict[str, dict]:
    variants = {}
    for phrase in phrases["phrases"]:
        for variant in phrase["variants"]:
            variants[str(variant["id"])] = {
                "id": str(variant["id"]),
                "korean": variant["korean"],
                "english": variant["english"],
                "male_ko": variant["speech"]["male"]["korean_pronunciation"],
                "female_ko": variant["speech"]["female"]["korean_pronunciation"],
            }
    return variants


def phrase_list(ids: list[str], variants: dict[str, dict]) -> str:
    if not ids:
        return "-"
    parts = []
    for phrase_id in ids:
        variant = variants.get(str(phrase_id))
        if variant:
            parts.append(f"{variant['id']}. {variant['korean']}")
        else:
            parts.append(str(phrase_id))
    return "\n".join(parts)


def phrase_pronunciation_list(ids: list[str], variants: dict[str, dict]) -> str:
    if not ids:
        return "-"
    parts = []
    for phrase_id in ids:
        variant = variants.get(str(phrase_id))
        if variant:
            parts.append(f"{variant['korean']}: 남 {variant['male_ko']} / 여 {variant['female_ko']}")
    return "\n".join(parts)


def keyword_index(assets: dict) -> dict[str, dict]:
    return {str(keyword["id"]): keyword for keyword in assets["keywords"]}


def keyword_list(ids: list[str], keywords: dict[str, dict]) -> str:
    if not ids:
        return "-"
    parts = []
    for keyword_id in ids:
        keyword = keywords.get(str(keyword_id))
        if keyword:
            pronunciation = keyword.get("koreanPronunciation") or keyword.get("korean_pronunciation")
            if pronunciation:
                parts.append(f"{keyword['korean']}({pronunciation})")
            else:
                parts.append(keyword["korean"])
        else:
            parts.append(str(keyword_id))
    return "\n".join(parts)


def daily_keyword_list(item: dict, keywords: dict[str, dict]) -> str:
    if item.get("lessonType") != "new":
        return "새 단어 없음\n(누적 복습)"
    return keyword_list(item.get("keyWordIds", []), keywords)


def item_by_id(schedule: list[dict]) -> dict[str, dict]:
    return {item["id"]: item for item in schedule}


def main() -> None:
    curriculum = load_json(CURRICULUM_PATH)
    phrases = load_json(PHRASES_PATH)
    assets = load_json(ASSETS_INDEX_PATH)
    variants = phrase_index(phrases)
    keywords = keyword_index(assets)
    schedule = curriculum["schedule"]
    weeks = {week["week"]: week for week in curriculum["weeks"]}

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Malgun Gothic"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Normal"].font.size = Pt(10)
    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("태국어 선교회화 25일 일별·주차별 학습 계획")
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4D78")
    title.paragraph_format.space_after = Pt(3)

    add_para(
        doc,
        "구성 원칙: 월~금 20일 동안 기본 문장을 새로 배우고, 마지막 5일은 전체 반복 복습으로 운영한다. "
        "토요일은 매주 6일차 오프라인 종합 복습 수업, 주일은 개인 활용 미션으로 둔다.",
        color="333333",
        size=10,
    )

    add_para(doc, "핵심 운영 요약", "Heading 1")
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    summary_table.rows[0].cells[0].text = "구분"
    summary_table.rows[0].cells[1].text = "내용"
    summary_rows = [
        ("새 문장 학습", "1~20일, 매일 새 문장 배치"),
        ("전체 반복 복습", "21~25일, 주제별 복습 후 25일 전체 리허설"),
        ("오프라인 수업", "각 주 토요일, 해당 주 학습 내용을 역할극으로 종합"),
        ("학습 항목 수", "기본 문장 전체. 세부 학습 ID는 좋은 아침/저녁, 예/아니오, 나이 질문/응답을 분리한다."),
        ("핵심 단어", "일별 학습 문장에서 새로 등장하는 단어를 중복 없이 배치하고 한국식 발음을 함께 표시"),
        ("자료 생성 기준", "카카오 알림 카드, PWA 학습 화면, 토요일 강의자료, 출력용 복습지를 이 계획에서 생성"),
    ]
    for label, detail in summary_rows:
        cells = summary_table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
    set_table_widths(summary_table, [1.5, 5.0])
    style_table(summary_table)

    add_para(doc, "주차별 계획", "Heading 1")
    week_table = doc.add_table(rows=1, cols=5)
    week_table.style = "Table Grid"
    headers = ["주차", "주제", "스토리", "토요일 수업", "주일 미션"]
    for idx, header in enumerate(headers):
        week_table.rows[0].cells[idx].text = header
    for week_no in range(1, 6):
        week = weeks[week_no]
        cells = week_table.add_row().cells
        cells[0].text = f"{week_no}주차"
        cells[1].text = week["theme"]
        cells[2].text = week["story"]
        cells[3].text = week["saturdayTitle"]
        cells[4].text = week["sundayTitle"]
    set_table_widths(week_table, [0.6, 1.25, 2.7, 1.05, 0.9])
    style_table(week_table)

    add_para(doc, "일별 학습 계획", "Heading 1")
    for week_no in range(1, 6):
        week = weeks[week_no]
        add_para(doc, f"{week_no}주차: {week['theme']}", "Heading 2")
        weekdays = [
            item
            for item in schedule
            if item.get("week") == week_no and item.get("dayType") == "weekday"
        ]
        daily_table = doc.add_table(rows=1, cols=7)
        daily_table.style = "Table Grid"
        for idx, header in enumerate(["일차", "유형", "제목", "새 문장", "핵심 단어", "복습 문장", "말하기 미션"]):
            daily_table.rows[0].cells[idx].text = header
        for item in weekdays:
            cells = daily_table.add_row().cells
            cells[0].text = str(item["day"])
            cells[1].text = "새 문장" if item.get("lessonType") == "new" else "복습"
            cells[2].text = item["title"]
            cells[3].text = phrase_list(item.get("newPhraseVariantIds", []), variants)
            cells[4].text = daily_keyword_list(item, keywords)
            cells[5].text = phrase_list(item.get("reviewVariantIds", []), variants)
            cells[6].text = item.get("speakingMission", "")
        set_table_widths(daily_table, [0.4, 0.55, 0.9, 1.15, 1.15, 1.25, 1.1])
        style_table(daily_table)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_para(doc, "토요일 오프라인 수업 계획", "Heading 1")
    sat_table = doc.add_table(rows=1, cols=5)
    sat_table.style = "Table Grid"
    for idx, header in enumerate(["주차", "수업 제목", "역할극 장면", "핵심 문장", "강사 진행 흐름"]):
        sat_table.rows[0].cells[idx].text = header
    for item in [i for i in schedule if i.get("dayType") == "saturday"]:
        cells = sat_table.add_row().cells
        cells[0].text = f"{item['week']}주차"
        cells[1].text = item["title"]
        cells[2].text = item["roleplayScene"]
        cells[3].text = phrase_list(item.get("roleplayFocusVariantIds", []), variants)
        cells[4].text = "\n".join(item.get("teacherFlow", []))
    set_table_widths(sat_table, [0.45, 1.0, 2.0, 1.55, 1.5])
    style_table(sat_table)

    add_para(doc, "주일 개인 복습 계획", "Heading 1")
    sun_table = doc.add_table(rows=1, cols=4)
    sun_table.style = "Table Grid"
    for idx, header in enumerate(["주차", "제목", "복습 문장", "개인 미션"]):
        sun_table.rows[0].cells[idx].text = header
    for item in [i for i in schedule if i.get("dayType") == "sunday"]:
        cells = sun_table.add_row().cells
        cells[0].text = f"{item['week']}주차"
        cells[1].text = item["title"]
        cells[2].text = phrase_list(item.get("reviewVariantIds", []), variants)
        cells[3].text = item.get("personalMission", "")
    set_table_widths(sun_table, [0.5, 1.25, 2.1, 2.65])
    style_table(sun_table)

    add_para(doc, "복습 주간 발음 참고", "Heading 1")
    add_para(
        doc,
        "마지막 5일의 복습 자료에는 학습자가 태국어 원문보다 한국식 발음을 먼저 보고 따라 할 수 있도록, "
        "문장 카드와 핵심 단어 카드에서 한국식 발음을 크게 노출한다.",
    )
    review_ids = ["1", "3", "4", "8", "9", "12", "13", "17", "19", "21", "22", "23", "24"]
    pronunciation_table = doc.add_table(rows=1, cols=2)
    pronunciation_table.style = "Table Grid"
    pronunciation_table.rows[0].cells[0].text = "대표 문장"
    pronunciation_table.rows[0].cells[1].text = "한국식 발음"
    for phrase_id in review_ids:
        variant = variants[phrase_id]
        cells = pronunciation_table.add_row().cells
        cells[0].text = f"{variant['id']}. {variant['korean']}"
        cells[1].text = f"남 {variant['male_ko']} / 여 {variant['female_ko']}"
    set_table_widths(pronunciation_table, [2.2, 4.3])
    style_table(pronunciation_table)

    footer = section.footer.paragraphs[0]
    footer.text = "Thai 25-day mission learning plan"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("666666")

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
