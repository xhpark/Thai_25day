#!/usr/bin/env python3
from __future__ import annotations

import json
import pathlib
from collections import defaultdict
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = pathlib.Path(__file__).resolve().parents[1]
PWA_DIR = ROOT / "assets/generated/pwa"
PPT_DIR = ROOT / "assets/generated/ppt"
MANIFEST_PATH = ROOT / "data/thai_image_source_manifest.json"
DOCS_DIR = ROOT / "docs"
OUT_DOCX = DOCS_DIR / "thai_25day_image_sourcing_request.docx"
OUT_MD = DOCS_DIR / "thai_25day_image_sourcing_request.md"


def load_json(path: pathlib.Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def local_plan(primary_image: dict) -> dict | None:
    for item in primary_image.get("sourcePlan", []):
        if item.get("type") == "local_preferred":
            return item
    return None


def stock_queries(primary_image: dict) -> list[str]:
    for item in primary_image.get("sourcePlan", []):
        if item.get("type") == "stock_preferred":
            return item.get("queries", [])
    return []


def source_status(primary_image: dict) -> str:
    plan = local_plan(primary_image)
    if plan and plan.get("status") == "ready":
        return "준비됨"
    return "사진 필요"


def collect_day_specs() -> list[dict]:
    items: list[dict] = []
    for path in sorted(PWA_DIR.glob("w*d*.json")):
        spec = load_json(path)
        if spec.get("dayType") == "sunday":
            continue
        items.append(
            {
                "id": spec["id"],
                "type": "평일",
                "week": spec.get("week"),
                "day": spec.get("day"),
                "title": spec.get("title", ""),
                "story": spec.get("story", ""),
                "missionGoal": spec.get("missionGoal", ""),
                "primaryImage": spec.get("primaryImage", {}),
            }
        )
    items.sort(key=lambda item: item["day"] or 0)

    saturday_items: list[dict] = []
    for path in sorted(PPT_DIR.glob("w*sat.json")):
        spec = load_json(path)
        saturday_items.append(
            {
                "id": spec["id"],
                "type": "토요일",
                "week": spec.get("week"),
                "day": None,
                "title": spec.get("title", ""),
                "story": spec.get("roleplayScene", ""),
                "missionGoal": spec.get("missionGoal", ""),
                "primaryImage": spec.get("primaryImage", {}),
            }
        )
    saturday_items.sort(key=lambda item: item["week"] or 0)
    return items + saturday_items


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.color.rgb = RGBColor(47, 125, 90)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(10)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(header_cells[idx], header, bold=True, color="173F31")
        set_cell_shading(header_cells[idx], "DFEEE5")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph("")


def build_markdown(items: list[dict], scene_groups: dict[str, list[dict]]) -> str:
    lines = [
        "# 태국어 선교회화 25일 이미지 수급 요청서",
        "",
        f"작성일: {date.today().isoformat()}",
        "",
        "## 현재 상태",
        "",
        "- Day 1 첫 인사 장면은 AI fallback 이미지로 파일럿에 임시 사용하기로 했습니다.",
        "- 나머지는 먼저 로컬 사진이나 무료/라이선스 확인 가능한 스톡 사진을 찾고, 적합한 사진이 없을 때만 AI 이미지를 사용합니다.",
        "- 일요일 복습 카드는 기본적으로 해당 주 평일/토요일 이미지를 재사용합니다.",
        "",
        "## 장면 그룹별 요청",
        "",
    ]
    for scene_id, group in scene_groups.items():
        first = group[0]
        primary = first["primaryImage"]
        ids = ", ".join(item["id"] for item in group)
        queries = "; ".join(stock_queries(primary)[:3])
        local = local_plan(primary) or {}
        lines.extend(
            [
                f"### {scene_id or 'fallback_scene'}",
                f"- 사용 날짜: {ids}",
                f"- 필요한 장면: {primary.get('altTextKo', '')}",
                f"- 권장 파일명: {local.get('assetPath', 'assets/images/scenes/<scene_id>.jpg')}",
                f"- 현재 상태: {source_status(primary)}",
                f"- 검색어: {queries}",
                "",
            ]
        )

    lines.extend(["## 날짜별 체크리스트", ""])
    for item in items:
        primary = item["primaryImage"]
        local = local_plan(primary) or {}
        lines.extend(
            [
                f"### {item['id']} {item['title']}",
                f"- 구분: {item['type']}",
                f"- 필요한 장면: {primary.get('altTextKo', '')}",
                f"- 스토리: {item.get('story') or item.get('missionGoal')}",
                f"- 파일명: {local.get('assetPath', '')}",
                f"- 상태: {source_status(primary)}",
                "",
            ]
        )
    return "\n".join(lines)


def build_docx(items: list[dict], scene_groups: dict[str, list[dict]]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("태국어 선교회화 25일 이미지 수급 요청서")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(23, 63, 49)

    subtitle = doc.add_paragraph()
    subtitle.add_run(f"작성일: {date.today().isoformat()} | 용도: 평일 PWA/카카오 카드, 토요일 강의자료 이미지 수급").font.size = Pt(10)

    add_heading(doc, "현재 상태와 원칙", 1)
    add_bullet(doc, "Day 1 첫 인사 장면은 AI fallback 이미지로 파일럿에 임시 사용합니다.")
    add_bullet(doc, "이미지는 로컬 직접 준비 사진을 우선하고, 다음으로 라이선스 확인 가능한 스톡 사진을 사용합니다.")
    add_bullet(doc, "적합한 로컬/스톡 사진이 없을 때만 AI-generated realistic image를 fallback으로 사용합니다.")
    add_bullet(doc, "일요일 복습 카드는 별도 신규 이미지보다 해당 주 평일/토요일 이미지를 재사용하는 방향입니다.")

    add_heading(doc, "장면 그룹별로 구하면 되는 사진", 1)
    group_rows = []
    for scene_id, group in scene_groups.items():
        first = group[0]
        primary = first["primaryImage"]
        ids = ", ".join(item["id"] for item in group)
        queries = "\n".join(stock_queries(primary)[:3])
        local = local_plan(primary) or {}
        group_rows.append(
            [
                scene_id or "fallback",
                ids,
                primary.get("altTextKo", ""),
                local.get("assetPath", ""),
                source_status(primary),
                queries,
            ]
        )
    add_table(
        doc,
        ["장면 ID", "사용 날짜", "필요한 장면", "권장 파일명", "상태", "검색어"],
        group_rows,
    )

    add_heading(doc, "날짜별 체크리스트", 1)
    day_rows = []
    for item in items:
        primary = item["primaryImage"]
        local = local_plan(primary) or {}
        day_label = f"{item['day']}일차" if item["day"] else f"{item['week']}주 토요일"
        day_rows.append(
            [
                item["id"],
                day_label,
                item["title"],
                primary.get("altTextKo", ""),
                local.get("assetPath", ""),
                source_status(primary),
            ]
        )
    add_table(
        doc,
        ["ID", "구분", "제목", "필요한 장면", "권장 파일명", "상태"],
        day_rows,
    )

    add_heading(doc, "사진 선택 기준", 1)
    add_bullet(doc, "모바일 화면에서 한눈에 주제와 행동이 보여야 합니다.")
    add_bullet(doc, "밝고 자연스러운 표정, 낮은 배경 clutter, 선명한 피사체를 우선합니다.")
    add_bullet(doc, "워터마크, 과도한 종교 상징화, 어둡거나 대립적인 분위기, 글자가 크게 박힌 사진은 피합니다.")
    add_bullet(doc, "가로 4:3 또는 3:2 사진이 가장 쓰기 좋고, 최소 1600x1200 이상을 권장합니다.")

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)


def main() -> None:
    items = collect_day_specs()
    scene_groups: dict[str, list[dict]] = defaultdict(list)
    for item in items:
        primary = item["primaryImage"]
        scene_id = primary.get("sceneId") or item["id"]
        scene_groups[scene_id].append(item)

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    OUT_MD.write_text(build_markdown(items, scene_groups), encoding="utf-8")
    build_docx(items, scene_groups)
    print(OUT_MD)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
