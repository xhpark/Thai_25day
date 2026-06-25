#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PPT_DIR = ROOT / "assets/generated/ppt"

TIME_ANCHORS = ["0:00~1:00", "1:00~3:00", "3:00~5:00", "5:00~8:00", "8:00~9:00", "9:00~10:00"]

# Optional hand-written coaching notes keyed by variantId. Falls back to a
# generic note when a variant has no specific tip.
COACHING_NOTES: dict[str, str] = {
    "1": "처음 만난 사람에게 쓰는 첫 인사입니다. 와이(합장) 동작과 함께 천천히 5회 따라 말하게 하세요. 남녀 말끝(크랍/카)을 짚어 주세요.",
    "2-1": "아침에 도착했을 때 쓰는 인사입니다. '떤 차오'(아침) 발음이 어려우면 음절을 나눠 두 번 더 들려주세요.",
    "2-2": "저녁 예배나 모임이 끝난 뒤 쓰는 인사입니다. 아침 인사와 짝지어 비교하며 말하게 하면 기억에 더 잘 남습니다.",
    "3": "도움을 받은 즉시 쓰는 감사 표현입니다. '막'(많이)에 강세를 살짝 주면 감사의 마음이 더 잘 전달됩니다.",
    "4": "감사 인사에 대한 응답입니다. 정말 고마워요와 반드시 짝으로 연습시켜, 한 사람이 말하면 다른 사람이 바로 받아치게 하세요.",
    "5-1": "짧고 자신 있게 말하는 것이 핵심입니다. 고개를 끄덕이는 동작을 함께 연습하세요.",
    "5-2": "예와 짝지어 번갈아 말하게 하면 짧은 응답에 대한 자신감이 빨리 붙습니다.",
}
DEFAULT_COACHING_NOTE = "정확한 발음보다 자신 있게 끝까지 말하는 태도를 칭찬해 주세요."

IMAGE_NOTE_TEMPLATE = (
    "이 장면 이미지는 '{alt}' 상황을 보여줍니다. 학습자가 문장을 들을 때 같은 장면을 떠올리도록, "
    "오디오를 재생하기 전에 이미지를 먼저 30초 정도 보여주고 시작하세요."
)


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_run(run, size=11, bold=False, color: str | None = None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc: Document, text: str = "", size=11, bold=False, color: str | None = None,
             space_after=6, line_spacing=1.25) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    return p


def add_heading(doc: Document, number: str, title: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {title}")
    style_run(run, size=16, bold=True, color="173F31")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "2F7D5A")
    border.append(bottom)
    p.paragraph_format.element.get_or_add_pPr().append(border)


def context_header_table(doc: Document, header: dict) -> None:
    table = doc.add_table(rows=3, cols=2)
    table.autofit = True
    rows = [
        ("수업 목표", header["lessonGoal"]),
        ("학습자가 마친 것", header["learnersHaveCovered"]),
        ("성공의 모습", header["successLooksLike"]),
    ]
    for idx, (label, value) in enumerate(rows):
        label_cell, value_cell = table.rows[idx].cells
        label_cell.width = Inches(1.6)
        value_cell.width = Inches(4.9)
        set_cell_shading(label_cell, "E8F4ED")
        label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        lp = label_cell.paragraphs[0]
        lrun = lp.add_run(label)
        style_run(lrun, size=11, bold=True, color="173F31")
        vp = value_cell.paragraphs[0]
        vrun = vp.add_run(value)
        style_run(vrun, size=11)
        vp.paragraph_format.line_spacing = 1.25


def build(lesson_id: str) -> Path:
    script_spec = load_json(PPT_DIR / f"{lesson_id.lower()}_script.json")

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title = doc.add_paragraph()
    trun = title.add_run(f"{script_spec['week']}주차 토요일 강사 스크립트")
    style_run(trun, size=22, bold=True, color="173F31")
    sub = doc.add_paragraph()
    srun = sub.add_run(script_spec["title"])
    style_run(srun, size=16, bold=True, color="2F7D5A")
    sub.paragraph_format.space_after = Pt(14)

    add_para(doc, "이 스크립트는 처음 이 수업을 맡는 강사도 바로 진행할 수 있도록 작성되었습니다.", size=10, color="5D6964")
    context_header_table(doc, script_spec["contextHeader"])

    add_heading(doc, "1", "수업 목표")
    add_para(doc, script_spec["missionGoal"])
    add_para(doc, f"오늘의 장면: {script_spec['roleplayScene']}", color="5D6964")

    add_heading(doc, "2", "시작하는 말 (그대로 읽어도 됩니다)")
    add_para(
        doc,
        f"“안녕하세요! 오늘은 {script_spec['title']} 시간이에요. "
        "이번 주에 배운 표현을 다 같이 한 번 따라 말해보고, 짝과 함께 짧은 역할극을 해볼 거예요. "
        "발음이 완벽하지 않아도 괜찮으니 큰 소리로, 자신 있게 말해보세요.”",
    )

    add_heading(doc, "3", "문장별 코칭 노트")
    for phrase in script_spec["phrases"]:
        female = phrase["speech"]["female"]
        male = phrase["speech"]["male"]
        p = doc.add_paragraph()
        run = p.add_run(f"{phrase['korean']}  ")
        style_run(run, size=13, bold=True, color="173F31")
        run2 = p.add_run(f"{female['thai']} / {male['thai']}")
        style_run(run2, size=13, color="1769D2")
        p.paragraph_format.space_after = Pt(2)
        add_para(
            doc,
            f"여: {female['korean_pronunciation']} · 남: {male['korean_pronunciation']}",
            size=10, color="5D6964", space_after=2,
        )
        note = COACHING_NOTES.get(phrase["variantId"], DEFAULT_COACHING_NOTE)
        add_para(doc, f"코칭 포인트: {note}", size=11, space_after=10)

    add_heading(doc, "4", "남성형 / 여성형 어미 안내")
    add_para(
        doc,
        "모든 문장은 남성 학습자는 '크랍(ครับ)'으로, 여성 학습자는 '카(ค่ะ)'로 끝맺습니다. "
        "강사는 학습자의 성별에 맞는 어미를 쓰도록 짝 활동 중에 짧게 점검해 주세요. "
        "성별을 묻기 애매한 경우, 본인이 편한 쪽으로 선택하게 해도 괜찮습니다.",
    )

    add_heading(doc, "5", "이미지 활용 노트")
    add_para(doc, IMAGE_NOTE_TEMPLATE.format(alt=script_spec.get("imageNotes", {}).get("primaryImage", {}).get("altTextKo", script_spec["roleplayScene"])))
    keyword_images = script_spec.get("imageNotes", {}).get("keywordImages", [])
    for ki in keyword_images:
        add_para(
            doc,
            f"단어 이미지({ki['keywordId']}): 아직 준비되지 않았다면 음성과 손짓으로 의미를 전달하고, "
            "이미지가 추가되는 대로 카드를 보여주며 다시 한 번 복습하세요.",
            size=10, color="5D6964",
        )

    add_heading(doc, "6", "역할극 진행 안내 (총 10분)")
    for idx, step in enumerate(script_spec.get("teacherFlow", [])):
        time_label = TIME_ANCHORS[idx] if idx < len(TIME_ANCHORS) else ""
        p = doc.add_paragraph()
        run = p.add_run(f"[{time_label}] ")
        style_run(run, size=12, bold=True, color="2F7D5A")
        run2 = p.add_run(step)
        style_run(run2, size=12, bold=True)
        p.paragraph_format.space_after = Pt(4)
    add_para(
        doc,
        "역할극 핵심 표현: " + ", ".join(p["korean"] for p in script_spec["phrases"]),
        size=11, color="5D6964",
    )

    add_heading(doc, "7", "선교 가이드 요약")
    add_para(doc, script_spec["ministryGuide"])
    add_para(doc, f"성공의 모습: {script_spec['contextHeader']['successLooksLike']}", size=11, color="5D6964")

    add_heading(doc, "8", "마무리 인사")
    add_para(
        doc,
        "“오늘 다들 발음보다 먼저 입을 열어 주셔서 감사합니다. 눈을 맞추고 미소로 인사한 것, "
        "정말 잘하셨어요. 내일은 집합 수업이 없으니, 오늘 배운 표현을 가족이나 가까운 분께 꼭 한 번 말해보세요. "
        "다음 주에 또 만나요!”",
    )

    out_path = PPT_DIR / f"{lesson_id.lower()}_script.docx"
    doc.save(str(out_path))
    return out_path


def main() -> int:
    parser = argparse.ArgumentParser(description="Build a Word teacher script from a teacher_script_spec JSON.")
    parser.add_argument("lesson_id", nargs="?", default="w1sat", help="e.g. w1sat, w2sat")
    args = parser.parse_args()
    out_path = build(args.lesson_id)
    print(f"[OK] wrote {out_path.relative_to(ROOT)}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
